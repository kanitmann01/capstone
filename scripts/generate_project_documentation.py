from __future__ import annotations

"""
CLI script: generate a styled Word document from project Markdown sources.

Reads the documentation pages under ``docs/``, converts headings, lists,
inline code, and fenced code blocks into formatted DOCX elements, appends
every ``scripts/*.py`` file as a traceability appendix, and writes
``docs/project-documentation.docx``.
"""

from datetime import datetime, timezone  # Standard library: UTC-aware timestamps
import re  # Standard library: regular expressions
from pathlib import Path  # Standard library: filesystem path abstraction

from docx import Document  # Third-party: Word document creation
from docx.enum.section import WD_SECTION  # Third-party: section break types
from docx.enum.style import WD_STYLE_TYPE  # Third-party: style type enumeration
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Third-party: text alignment
from docx.oxml import OxmlElement  # Third-party: low-level XML element builder
from docx.oxml.ns import qn  # Third-party: XML namespace helper
from docx.shared import Inches  # Third-party: imperial unit conversion
from docx.shared import Pt  # Third-party: point unit conversion
from docx.shared import RGBColor  # Third-party: colour values


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = PROJECT_ROOT / "docs"
OUTPUT_FILE = DOCS_DIR / "project-documentation.docx"
SOURCE_FILES = [
    "00-comprehensive-project-handbook.md",
    "README.md",
    "01-project-overview.md",
    "02-repository-tour.md",
    "03-system-architecture.md",
    "04-api-reference.md",
    "05-scanner-modules.md",
    "06-threat-intelligence-and-caching.md",
    "07-web-ui-and-user-experience.md",
    "08-baseline-evaluation-and-data.md",
    "09-development-testing-and-operations.md",
    "10-agency-agents-subtree.md",
    "11-risks-limitations-and-roadmap.md",
    "12-ml-lab-and-classifier.md",
    "13-brand-login-capstone-methodology.md",
    "../old-data/14-brand-login-eda.md",
]

TITLE_COLOR = RGBColor(15, 23, 42)
ACCENT_COLOR = RGBColor(14, 116, 144)
BODY_COLOR = RGBColor(31, 41, 55)
MUTED_COLOR = RGBColor(71, 85, 105)
CODE_SHADE = "EEF4F7"
LIST_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def add_field(paragraph, field_code: str) -> None:
    """Insert a Word field (e.g., PAGE) into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = " "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def set_cell_shading(style, fill: str) -> None:
    """Apply a background colour to a paragraph style."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    style.element.get_or_add_pPr().append(shading)


def configure_styles(document: Document) -> None:
    """Set up document-wide font, colour, and spacing styles."""
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = TITLE_COLOR

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(11)
    subtitle.font.italic = True
    subtitle.font.color.rgb = MUTED_COLOR

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11.5)):
        heading = styles[name]
        heading.font.name = "Aptos Display"
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = TITLE_COLOR if name == "Heading 1" else ACCENT_COLOR

    if "DocCode" not in styles:
        code_style = styles.add_style("DocCode", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["DocCode"]
    code_style.base_style = styles["Normal"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = BODY_COLOR
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.right_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    set_cell_shading(code_style, CODE_SHADE)

    if "DocSmall" not in styles:
        small_style = styles.add_style("DocSmall", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small_style = styles["DocSmall"]
    small_style.base_style = styles["Normal"]
    small_style.font.name = "Aptos"
    small_style.font.size = Pt(9)
    small_style.font.color.rgb = MUTED_COLOR


def configure_page(document: Document) -> None:
    """Configure page margins and footer page-number fields."""
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style = document.styles["DocSmall"]
    footer_para.add_run("Project Documentation | Page ")
    add_field(footer_para, "PAGE")


def add_cover(document: Document) -> None:
    """Render the title page with project name and generation timestamp."""
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Capstone Project Documentation")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Comprehensive mixed-audience reference for the phishing scanner "
        "application and the wider repository."
    )

    meta = document.add_paragraph(style="DocSmall")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta.add_run(f"Generated from repository Markdown sources on {generated_at}")

    document.add_paragraph("")

    summary = document.add_paragraph(style="Normal")
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        "This document consolidates the authored documentation pages under "
        "`docs/` into a styled handoff artifact."
    )

    document.add_section(WD_SECTION.NEW_PAGE)


def add_toc(document: Document) -> None:
    """Insert an auto-generated Table of Contents field."""
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Table of Contents")

    toc_paragraph = document.add_paragraph(style="Normal")
    add_field(toc_paragraph, r'TOC \o "1-3" \h \z \u')

    note = document.add_paragraph(style="DocSmall")
    note.add_run("If the table of contents appears empty, update fields in Word.")

    document.add_section(WD_SECTION.NEW_PAGE)


def add_inline_runs(paragraph, text: str) -> None:
    """Split inline backtick code and add styled runs to a paragraph."""
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def emit_paragraph(document: Document, text: str) -> None:
    """Route a single Markdown line into the appropriate DOCX paragraph style."""
    stripped = text.strip()
    if not stripped:
        document.add_paragraph("")
        return

    if stripped.startswith("# "):
        p = document.add_paragraph(style="Heading 1")
        add_inline_runs(p, stripped[2:].strip())
        return

    if stripped.startswith("## "):
        p = document.add_paragraph(style="Heading 2")
        add_inline_runs(p, stripped[3:].strip())
        return

    if stripped.startswith("### "):
        p = document.add_paragraph(style="Heading 3")
        add_inline_runs(p, stripped[4:].strip())
        return

    if stripped.startswith("- "):
        p = document.add_paragraph(style="List Bullet")
        add_inline_runs(p, stripped[2:].strip())
        return

    ordered = LIST_RE.match(stripped)
    if ordered:
        p = document.add_paragraph(style="List Number")
        add_inline_runs(p, ordered.group(2).strip())
        return

    p = document.add_paragraph(style="Normal")
    add_inline_runs(p, text)


def add_markdown_file(document: Document, path: Path) -> None:
    """Parse a Markdown file and emit corresponding DOCX elements."""
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                for code_line in code_buffer:
                    document.add_paragraph(code_line, style="DocCode")
                if not code_buffer:
                    document.add_paragraph("", style="DocCode")
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        emit_paragraph(document, line)

    if code_buffer:
        for code_line in code_buffer:
            document.add_paragraph(code_line, style="DocCode")


def add_scripts_appendix(document: Document) -> None:
    """Append every ``scripts/*.py`` file as traceable source (same folder as this generator)."""
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Automation scripts (scripts/)")

    intro = document.add_paragraph(style="Normal")
    intro.add_run(
        "The sections below embed the current contents of each Python file under "
        "`scripts/`, so the Word export stays aligned with the runnable tooling in "
        "that directory."
    )

    scripts_dir = PROJECT_ROOT / "scripts"
    py_files = sorted(scripts_dir.glob("*.py"))
    if not py_files:
        note = document.add_paragraph(style="DocSmall")
        note.add_run("No Python files found in scripts/.")
        return

    for path_index, path in enumerate(py_files):
        if path_index:
            document.add_paragraph("")

        title = document.add_paragraph(style="Heading 2")
        title.add_run(path.name)

        rel = path.relative_to(PROJECT_ROOT)
        sub = document.add_paragraph(style="DocSmall")
        sub.add_run(str(rel).replace("\\", "/"))

        body = path.read_text(encoding="utf-8", errors="replace").splitlines()
        if not body:
            document.add_paragraph("(empty file)", style="DocCode")
            continue
        for line in body:
            document.add_paragraph(line, style="DocCode")


def build_document() -> Path:
    """Assemble the full DOCX from all Markdown sources."""
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)
    add_toc(document)

    for index, filename in enumerate(SOURCE_FILES):
        add_markdown_file(document, DOCS_DIR / filename)
        if index < len(SOURCE_FILES) - 1:
            document.add_section(WD_SECTION.NEW_PAGE)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_scripts_appendix(document)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


def main() -> int:
    """Entry point: build and save the project documentation Word file."""
    output = build_document()
    print(f"Wrote {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
